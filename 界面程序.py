import tkinter as tk
from tkinter import messagebox
import execjs
import csv
import datetime
from time import sleep
from snownlp import SnowNLP
import requests
import json
import re
from docx import Document
from docx.shared import Inches
import os
import glob
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import numpy as np
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import random
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
import pandas as pd

#删除冗余文件
for pattern in ['fu_*.txt', 'screenshot_*.png', 'note_*.xlsx', 'xhs_*.xlsx','*.docx','*.csv']:
            for file in glob.glob(pattern):
                os.remove(file)
#获取cookie文件
urls = [
    "https://www.baidu.com/link?url=aXGalUXyzrCO7M510Bb8iZjfP_NghcPrkOnY0ShyQML-vGRJBKjOLh89cRjy6VEH&wd=&eqid=84db72770083d32a0000000665e53abb",
    "https://www.baidu.com/link?url=l64UVulqbFZLlFyjD6P3aV_TOE3_OWB4bW-1u-mR-O8rPmvA0EAgHaCMq571Murw&wd=&eqid=9f682c5600f435510000000665e53ade",
    "https://www.baidu.com/link?url=uG3sW0yfRXTIRtat1kpA_f_WrlsPpvs5Cf_Csp8tmRAXUe75cyhNXEtrmawMw1_u&wd=&eqid=de96f5d000001f6f0000000665dac3ce",
    "https://www.baidu.com/link?url=lGqYCGniL7PtJ1DREzCBQqVd07s3KFBsXiESdnNAqp62A53R7JEkj8pZz3kP724y&wd=&eqid=8267aa0a00f451f30000000665e53af1",
    "https://www.baidu.com/link?url=0rlK9EGZkAHO9HXFOdOApBQdOI3ZRKygKcPM6z-zevawjl8VebOhQXBCmB2ScDpl&wd=&eqid=b526dbc600f410120000000665e53afb",
    "https://www.baidu.com/link?url=l5rZsPOhXfKmYsI0CCFAItwxlqpbM5lVbRh5nYvfEStp0V2lzv8xELN1OG75Zi_O&wd=&eqid=c143e9c800f3e7e40000000665e53b05",
    "https://www.baidu.com/link?url=nkzhVr3uSv1-_bl28p0VA8wFDo1fF15_sCdz9XJUde8Br6rP5segHUtAB1at7NeY&wd=&eqid=d862a6780081babd0000000665e53b0c",
    "https://www.baidu.com/link?url=PJGKS_bpaBSL3yHFLEZDh-JvKnQpqtuM7SS6rir2jwqQh82B9TWCW1wpaYCu2PeL&wd=&eqid=863ec96b0080c1c90000000665e53b21",
    "https://www.baidu.com/link?url=UQviTRc-xlevI6q0rTTEJGNhMKHMduJuif9koWN4hQLPOaxAqnnuI1woSspji1b8&wd=&eqid=8f6ccdc300f754e10000000665e53b28",
    "https://www.baidu.com/link?url=hNcf_dpLvZPkgERKCBV0prATWqE7rrMU9a4TmGkO3gzgpOjKxVBCOuoGx6gjnqnA&wd=&eqid=c037ff1800f56df20000000665e53b2f",
    "https://www.baidu.com/link?url=P4GFZ4uvQlEHDGp1mEtT2wcH_WKUQrRmXrJUpslCr-uRZV5OEC2ADbkWS_pI8TZW&wd=&eqid=e93e65ae0011e2d70000000665e53b3f"
]
#获取cookie文件
if not os.path.exists('cookies.txt'):
    options = Options()
    options.add_argument("--disable-blink-features=AutomationControlled")
    driver = webdriver.Chrome(options=options)
    url = random.choice(urls)
    driver.get(url)
    time.sleep(30)
    cookie_list = driver.get_cookies()
    cookies = {}
    for i in cookie_list:
        key = i['name']
        value = i['value']
        cookies[key] = value
    with open('cookies.txt', 'w') as f:
        f.write(str(cookies))
    driver.close()
else:
    pass
#第一页
class PageOne(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller

        label = tk.Label(self, text='''
        用前必看：
        仅用于学习用途，侵权必删；
        每个邮箱添加的时候需要点击添加邮箱按钮,对于添加的新的邮箱，空白会默认跳过；
        每个时间添加的时候需要点击添加时间按钮,对于添加的新的时间，空白会默认跳过；
        输入时间格式为小时:分钟，其中的“:”是英文格式的，不是中文格式的；
        检查完后点击"完成"按钮后才默认开始运行进程
        点击"完成"按钮后界面会卡住但请不要关闭界面；
        中途可能会跳出浏览器弹窗，是正常现象，这是在给页面截图，不需要额外操作；
        谢谢使用。
        如继续使用请点击下一步。''')
        label.pack(pady=10, padx=10)

        button = tk.Button(self, text="下一步", command=lambda: controller.show_page(PageTwo))
        button.pack(pady=10)
#第二页
class PageTwo(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller
        self.keywords = tk.StringVar()
        self.pages = tk.StringVar()
        self.emails =[]

        label = tk.Label(self, text="请输入相对应值")
        label.pack(pady=10, padx=10)

        keyword_label = tk.Label(self, text="关键词(多个关键词用空格隔开):")
        keyword_label.pack()
        keyword_entry = tk.Entry(self, textvariable=self.keywords)
        keyword_entry.pack()

        pages_label = tk.Label(self, text="爬取的页数(无特殊需求不建议设置太多页数):")
        pages_label.pack()
        pages_entry = tk.Entry(self, textvariable=self.pages)
        pages_entry.pack()

        email_label = tk.Label(self, text="收取报告的邮箱地址:")
        email_label.pack()

        self.email_frame = tk.Frame(self)  # 创建一个新的框架用于显示邮箱地址
        self.email_frame.pack()

        self.add_email_entry()  # 默认显示一个输入框

        add_email_button = tk.Button(self, text="添加邮箱", command=self.add_email)
        add_email_button.pack(pady=10)

        prev_button = tk.Button(self, text="上一步", command=lambda: controller.show_page(PageOne))
        prev_button.pack(pady=10)

        self.next_button = tk.Button(self, text="下一步", command=lambda: controller.show_page(PageThree))
        self.next_button.pack(pady=10)
        self.next_button.configure(state="disabled")

        finish_button = tk.Button(self, text="检查", command=self.finish)
        finish_button.pack(pady=10)

    def add_email_entry(self):
        email_entry = tk.Entry(self.email_frame)
        email_entry.pack(side="left", padx=5)
        self.emails.append(email_entry)

    def add_email(self):
        email = self.emails[-1].get()
        if email.strip():  # 跳过空白值
            self.add_email_entry()
            self.emails[-1].delete(0, tk.END)
    #检查是否出错
    def finish(self):
        if not self.keywords.get():
            messagebox.showerror("错误", "请至少输入一个关键词")
            return

        if not self.pages.get():
            messagebox.showerror("错误", "请至少输入页数")
            return

        try:
            page_count = int(self.pages.get())

            if page_count < 2:
                messagebox.showerror("错误", "至少输入2页")
                return
        except ValueError:
            messagebox.showerror("错误", "页数必须为整数")
            return

        if not self.emails:
            messagebox.showerror("错误", "请至少输入一个邮箱")
            return

        # 检查邮箱格式
        email_list = []
        for email_entry in self.emails:
            email = email_entry.get()
            if email.strip():  # 跳过空白值
                if not self.check_email_format(email):
                    messagebox.showerror("错误", "邮箱输入错误")
                    return
                email_list.append(email)
        if not email_list:
            messagebox.showerror("错误", "请至少输入一个邮箱")
            return

        # 执行相关操作，输出值
        self.emails = [email_entry.get() for email_entry in self.emails]
        keyword_list = self.keywords.get().split()
        email_list = self.emails
        messagebox.showinfo("结果是否已经确定好？", f"关键词: {', '.join(keyword_list)}\n页数: {page_count}\n邮箱: {', '.join(email_list)}")

        self.next_button.configure(state="normal")

    def check_email_format(self, email):
        # 简单的邮箱格式检查
        if "@" in email and "." in email:
            return True
        return False

    def get_emails(self):
        return self.emails
#第三页
class PageThree(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller

        self.times = []

        label = tk.Label(self, text="请输入运行时间(小时:分钟，中间的符号是英文符号)")
        label.pack(pady=10, padx=10)

        self.time_frame = tk.Frame(self)
        self.time_frame.pack()

        self.add_time_entry()

        add_time_button = tk.Button(self, text="添加时间", command=self.add_time)
        add_time_button.pack(pady=10)

        prev_button = tk.Button(self, text="上一步", command=lambda: controller.show_page(PageTwo))
        prev_button.pack(pady=10)

        check_button = tk.Button(self, text="检查", command=self.check)
        check_button.pack(pady=10)

        self.finish_button = tk.Button(self, text="完成", command=self.start_countdown)
        self.finish_button.pack(pady=10)
        self.finish_button.configure(state="disabled")

        self.is_running = False

    def add_time_entry(self):
        time_entry = tk.Entry(self.time_frame)
        time_entry.pack(side="left", padx=5)
        self.times.append(time_entry)

    def add_time(self):
        self.add_time_entry()

    def check(self):
        time_list = [time_entry.get() for time_entry in self.times if time_entry.get().strip()]
        if not time_list:
            messagebox.showerror("错误", "请至少输入一个时间")
            return

        try:
            current_time = datetime.datetime.now().time()
            current_datetime = datetime.datetime.combine(datetime.date.today(), current_time)
            target_datetimes = []

            for time_str in time_list:
                time_obj = datetime.datetime.strptime(time_str, "%H:%M")
                target_datetime = datetime.datetime.combine(datetime.date.today(), time_obj.time())

                if current_datetime > target_datetime:
                    target_datetime += datetime.timedelta(days=1)

                target_datetimes.append(target_datetime)

            time_differences = [target_datetime - current_datetime for target_datetime in target_datetimes]
            time_differences_str = [str(time_difference).split('.')[0] for time_difference in time_differences]
            messagebox.showinfo("结果是否已经确定好？", f"程序距离最近的运行时间还有：{', '.join(time_differences_str)}")
            self.finish_button.configure(state="normal")
        except ValueError:
            messagebox.showerror("错误", "时间格式无效")

    def start_countdown(self):
        time_list = [time_entry.get() for time_entry in self.times if time_entry.get().strip()]
        if not time_list:
            messagebox.showerror("错误", "请至少输入一个时间")
            return

        print(f"运行时间: {', '.join(time_list)}")
        self.finish_button.configure(state="disabled")

        self.is_running = True
        self.run_programs(time_list)

    def run_programs(self, time_list):
        current_time = datetime.datetime.now().time()
        current_datetime = datetime.datetime.combine(datetime.date.today(), current_time)
        target_datetimes = []

        for time_str in time_list:
            time_obj = datetime.datetime.strptime(time_str, "%H:%M")
            target_datetime = datetime.datetime.combine(datetime.date.today(), time_obj.time())

            if current_datetime > target_datetime:
                target_datetime += datetime.timedelta(days=1)

            target_datetimes.append(target_datetime)

        while self.is_running:
            closest_time = min(target_datetimes)
            time_difference = closest_time - current_datetime
            time_difference_str = str(time_difference).split('.')[0]

            if time_difference.total_seconds() <= 0:
                break

            print(f"距离最近的运行时间还有：{time_difference_str}")

            time.sleep(1)
            current_datetime = datetime.datetime.now()

        self.run_spider()
        self.run_spider_1()
        self.process_xiaohongshu_data()
        self.send_email_with_attachments()

        if self.is_running:
            self.start_countdown()

    def finish(self):
        time_list = [time_entry.get() for time_entry in self.times if time_entry.get().strip()]
        if not time_list:
            messagebox.showerror("错误", "请至少输入一个时间")
            return
        print(f"运行时间: {', '.join(time_list)}")
        self.run_spider()
        self.run_spider_1()
        self.process_xiaohongshu_data()
        self.send_email_with_attachments()

    note_count = 0
    #爬取话题笔记
    def run_spider(self):
        global note_count
        header = ["笔记关键词", "笔记标题", "小红书账号名", "笔记发布时间",
                  "笔记收藏数量", "笔记评论数量", "笔记点赞数量", "笔记转发数量", "笔记内容", "笔记链接"]
        f = open(f"话题笔记.csv", "a", encoding="utf-8-sig", newline="")
        writer = csv.DictWriter(f, header)
        writer.writeheader()
        user_agents = [
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36 Edg/115.0.1901.188",
            "Mozilla/5.0 (Windows NT 10.0; WOW64; rv:54.0) Gecko/20100101 Firefox/54.0",
            "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.94 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.84 Safari/537.36",
            "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.26 Safari/537.36 Core/1.63.6788.400 QQBrowser/10.3.2727.400",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3",
            "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:54.0) Gecko/20100101 Firefox/54.0",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_6) AppleWebKit/603.3.8 (KHTML, like Gecko) Version/10.1.2 Safari/603.3.8",
            "Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/60.0.3112.90 Safari/537.36",
            "Mozilla/5.0 (iPad; CPU OS 10_3_3 like Mac OS X) AppleWebKit/603.1.30 (KHTML, like Gecko) Version/10.0 Mobile/14G60 Safari/602.1"
        ]

        headers = {
            "authority": "edith.xiaohongshu.com",
            "accept": "application/json, text/plain, */*",
            "accept-language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
            "content-type": "application/json;charset=UTF-8",
            "origin": "https://www.xiaohongshu.com",
            "referer": "https://www.xiaohongshu.com/",
            "user-agent": random.choice(user_agents)
        }
        with open('cookies.txt', 'r') as f:
            cookies = eval(f.read())
        js = execjs.compile(open(r'info.js', 'r', encoding='utf-8').read())
        note_count = 0

        def get_time(ctime):
            timeArray = time.localtime(int(ctime / 1000))
            otherStyleTime = time.strftime("%Y.%m.%d", timeArray)
            return str(otherStyleTime)

        def get_note_info(note_id):
            note_url = 'https://edith.xiaohongshu.com/api/sns/web/v1/feed'
            data = {
                "source_note_id": note_id,
                "image_scenes": [
                    "CRD_PRV_WEBP",
                    "CRD_WM_WEBP"
                ]
            }
            data = json.dumps(data, separators=(',', ':'))
            ret = js.call('get_xs', '/api/sns/web/v1/feed', data, cookies['a1'])
            headers['x-s'], headers['x-t'] = ret['X-s'], str(ret['X-t'])
            response = requests.post(note_url, headers=headers, cookies=cookies, data=data)
            json_data = response.json()
            try:
                note_data = json_data['data']['items'][0]
            except:
                print(f'笔记 {note_id} 不允许查看')
                return
            sava_data(note_data)

        def keyword_search(keyword, page_count):
            api = '/api/sns/web/v1/search/notes'
            search_url = "https://edith.xiaohongshu.com/api/sns/web/v1/search/notes"
            data = {
                "image_scenes": "FD_PRV_WEBP,FD_WM_WEBP",
                "keyword": "",
                "note_type": "0",
                "page": "",
                "page_size": "20",
                "search_id": "2c7hu5b3kzoivkh848hp0",
                "sort": "general"
            }
            data = json.dumps(data, separators=(',', ':'))
            data = re.sub(r'"keyword":".*?"', f'"keyword":"{keyword}"', data)

            for page in range(1, page_count):
                data = re.sub(r'"page":".*?"', f'"page":"{page}"', data)
                ret = js.call('get_xs', api, data, cookies['a1'])
                headers['x-s'], headers['x-t'] = ret['X-s'], str(ret['X-t'])
                response = requests.post(search_url, headers=headers, cookies=cookies, data=data.encode('utf-8'))
                json_data = response.json()
                try:
                    notes = json_data['data']['items']
                except:
                    break
                for note in notes:
                    note_id = note['id']
                    if len(note_id) != 24:
                        continue
                    get_note_info(note_id)
                    time.sleep(random.uniform(3, 6))

        def sava_data(note_data):
            global note_count
            try:
                ip_location = note_data['note_card']['ip_location']
            except:
                ip_location = '未知'
            if get_time(note_data['note_card']['time']) > date_time:
                return 0
            note_link = "https://www.xiaohongshu.com/discovery/item/" + note_data['id'] if 'id' in note_data else '无'
            data_dict = {
                "笔记关键词": keyword.strip(),
                "笔记标题": note_data['note_card']['title'].strip(),
                "小红书账号名": note_data['note_card']['user']['nickname'].strip(),
                "笔记发布时间": get_time(note_data['note_card']['time']),
                "笔记收藏数量": note_data['note_card']['interact_info']['collected_count'],
                "笔记评论数量": note_data['note_card']['interact_info']['comment_count'],
                "笔记点赞数量": note_data['note_card']['interact_info']['liked_count'],
                "笔记转发数量": note_data['note_card']['interact_info']['share_count'],
                "笔记内容": note_data['note_card']['desc'].strip().replace('\n', ''),
                "笔记链接": note_link
            }
            global note_count
            note_count += 1
            print(f"当前笔记关键词: {keyword.strip()}\n",
                  f"当前笔记数量: {note_count}\n",
                  f"笔记标题：{data_dict['笔记标题']}\n",
                  f"小红书账号名：{data_dict['小红书账号名']}\n",
                  f"笔记发布时间：{data_dict['笔记发布时间']}\n",
                  f"笔记收藏数量：{data_dict['笔记收藏数量']}\n",
                  f"笔记评论数量：{data_dict['笔记评论数量']}\n",
                  f"笔记点赞数量：{data_dict['笔记点赞数量']}\n",
                  f"笔记转发数量：{data_dict['笔记转发数量']}\n",
                  f"笔记内容：{data_dict['笔记内容']}\n",
                  f"笔记链接：{data_dict['笔记链接']}\n"
                  )
            writer.writerow(data_dict)

        date_time = time.strftime('%Y.%m.%d', time.localtime())
        page_count = int(self.controller.pages[PageTwo].pages.get())
        keyword = self.controller.pages[PageTwo].keywords.get()
        keyword_list = keyword.split(" ")
        for keyword in keyword_list:
            keyword_search(keyword.strip(), page_count)
    #爬取话题笔记下的评论

    def run_spider_1(self):
        user_agents = [
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36 Edg/115.0.1901.188",
            "Mozilla/5.0 (Windows NT 10.0; WOW64; rv:54.0) Gecko/20100101 Firefox/54.0",
            "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.94 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.84 Safari/537.36",
            "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.26 Safari/537.36 Core/1.63.6788.400 QQBrowser/10.3.2727.400",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3",
            "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:54.0) Gecko/20100101 Firefox/54.0",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_6) AppleWebKit/603.3.8 (KHTML, like Gecko) Version/10.1.2 Safari/603.3.8",
            "Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/60.0.3112.90 Safari/537.36",
            "Mozilla/5.0 (iPad; CPU OS 10_3_3 like Mac OS X) AppleWebKit/603.1.30 (KHTML, like Gecko) Version/10.0 Mobile/14G60 Safari/602.1"
        ]

        headers = {
            "authority": "edith.xiaohongshu.com",
            "accept": "application/json, text/plain, */*",
            "accept-language": "zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6",
            "content-type": "application/json;charset=UTF-8",
            "origin": "https://www.xiaohongshu.com",
            "referer": "https://www.xiaohongshu.com/",
            "user-agent": random.choice(user_agents)
        }

        with open('cookies.txt', 'r') as f:
            cookies = eval(f.read())

        def trans_date(v_timestamp):
            v_timestamp = int(str(v_timestamp)[:10])
            timeArray = time.localtime(v_timestamp)
            otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
            return otherStyleTime

        df = pd.read_csv('话题笔记.csv')
        if '笔记链接' in df.columns and not df['笔记链接'].isnull().all():
            df['note_id'] = df['笔记链接'].apply(lambda x: x.replace('https://www.xiaohongshu.com/discovery/item/', ''))
            note_id_list = df['note_id'].tolist()
            now = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
            result_file = '小红书评论_{}.csv'.format(now)
        else:
            pass

        for note_id in note_id_list:
            page = 1
            while True:
                if page == 1:
                    url = 'https://edith.xiaohongshu.com/api/sns/web/v2/comment/page?note_id={}&top_comment_id=&image_scenes=FD_WM_WEBP,CRD_WM_WEBP'.format(
                        note_id)
                else:
                    url = 'https://edith.xiaohongshu.com/api/sns/web/v2/comment/page?note_id={}&top_comment_id=&image_scenes=FD_WM_WEBP,CRD_WM_WEBP&cursor={}'.format(
                        note_id, next_cursor)
                r = requests.get(url, headers=headers, cookies=cookies)
                if str(r.status_code) != '200':
                    continue
                json_data = r.json()
                # 定义一些空列表用于存放数据
                content_list = []  # 评论内容
                create_time_list = []  # 评论时间
                ip_list = []  # 评论IP属地
                like_count_list = []  # 评论点赞数
                nickname_list = []  # 评论者昵称
                user_id_list = []  # 评论者id
                user_link_list = []  # 评论者主页链接
                comment_level_list = []  # 评论级别
                try:
                    comment_num = len(json_data['data']['comments'])
                except KeyError:
                    comment_num = 0
                #comment_num = len(json_data['data']['comments'])
                print('开始爬【{}】第{}页，笔记数量：{}'.format(note_id, page, comment_num))
                sleep(random.uniform(1, 3))
                # 循环解析json数据
                print('开始爬取【一级】评论')
                for c in json_data['data']['comments']:
                    # 评论者昵称
                    nickname = c['user_info']['nickname']
                    nickname_list.append(nickname)
                    # 评论者id
                    user_id = c['user_info']['user_id']
                    user_id_list.append(user_id)
                    # 评论者主页链接
                    user_link = 'https://www.xiaohongshu.com/user/profile/' + user_id
                    user_link_list.append(user_link)
                    # 评论时间
                    create_time = trans_date(c['create_time'])
                    create_time_list.append(create_time)
                    # 评论IP属地
                    try:
                        ip = c['ip_location']
                    except:
                        ip = ''
                    ip_list.append(ip)
                    # 评论点赞数
                    like_count = c['like_count']
                    like_count_list.append(like_count)
                    # 评论内容
                    content = c['content']
                    content_list.append(content)
                    # 二级评论
                    sub_comment_count = c['sub_comment_count']
                    # 根评论id
                    root_comment_id = c['id']
                    # 评论级别
                    comment_level_list.append('根评论')
                    if int(sub_comment_count) > 0:
                        print('开始爬取【二级】评论')
                        for c2 in c['sub_comments']:
                            # 评论者昵称
                            nickname2 = c2['user_info']['nickname']
                            nickname_list.append(nickname2)
                            # 评论者id
                            user_id2 = c2['user_info']['user_id']
                            user_id_list.append(user_id2)
                            # 评论者主页链接
                            user_link2 = 'https://www.xiaohongshu.com/user/profile/' + user_id2
                            user_link_list.append(user_link2)
                            # 评论时间
                            create_time2 = trans_date(c2['create_time'])
                            create_time_list.append(create_time2)
                            # 评论IP属地
                            try:
                                ip2 = c2['ip_location']
                            except:
                                ip2 = ''
                            ip_list.append(ip2)
                            # 评论点赞数
                            like_count2 = c2['like_count']
                            like_count_list.append(like_count2)
                            # 评论内容
                            content2 = c2['content']
                            content_list.append(content2)
                            # 评论级别
                            comment_level_list.append('二级评论')
                    # 展开评论
                    if c['sub_comment_has_more'] == True:
                        print('开始爬取【二级展开】评论')
                        extend_page = 1
                        while True:
                            if extend_page == 1:
                                url_more = 'https://edith.xiaohongshu.com/api/sns/web/v2/comment/sub/page?note_id={}&root_comment_id={}&image_scenes=FD_WM_WEBP,CRD_WM_WEBP&cursor={}&num=10'.format(
                                    note_id, root_comment_id, c['sub_comment_cursor'])
                            else:
                                url_more = 'https://edith.xiaohongshu.com/api/sns/web/v2/comment/sub/page?note_id={}&root_comment_id={}&image_scenes=FD_WM_WEBP,CRD_WM_WEBP&cursor={}&num=10'.format(
                                    note_id, root_comment_id, next_cursor_more)
                            r_more = requests.get(url_more, headers=headers, cookies=cookies)
                            json_data_more = r_more.json()
                            # 循环解析json数据
                            for c_more in json_data_more['data']['comments']:
                                # 评论者昵称
                                nickname_more = c_more['user_info']['nickname']
                                nickname_list.append(nickname_more)
                                # 评论者id
                                user_id_more = c_more['user_info']['user_id']
                                user_id_list.append(user_id_more)
                                # 评论者主页链接
                                user_link_more = 'https://www.xiaohongshu.com/user/profile/' + user_id_more
                                user_link_list.append(user_link_more)
                                # 评论时间
                                create_time_more = trans_date(c_more['create_time'])
                                create_time_list.append(create_time_more)
                                # 评论IP属地
                                try:
                                    ip_more = c_more['ip_location']
                                except:
                                    ip_more = ''
                                ip_list.append(ip_more)
                                # 评论点赞数
                                like_count_more = c_more['like_count']
                                like_count_list.append(like_count_more)
                                # 评论内容
                                content_more = c_more['content']
                                content_list.append(content_more)
                                # 评论级别
                                comment_level_list.append('二级展开评论')
                            if not json_data_more['data']['has_more']:
                                break
                            # 判断终止条件
                            next_cursor_more = json_data_more['data']['cursor']
                            extend_page += 1
                # 保存数据到DF
                df = pd.DataFrame(
                    {
                        '笔记链接': 'https://www.xiaohongshu.com/explore/' + note_id,
                        '页码': page,
                        '评论者昵称': nickname_list,
                        '评论者id': user_id_list,
                        '评论者主页链接': user_link_list,
                        '评论时间': create_time_list,
                        '评论IP属地': ip_list,
                        '评论点赞数': like_count_list,
                        '评论级别': comment_level_list,
                        '评论内容': content_list,
                    }
                )
                # 设置csv文件表头
                if os.path.exists(result_file):
                    header = False
                else:
                    header = True
                # 保存到csv
                df.to_csv(result_file, mode='a+', header=header, index=False, encoding='utf_8_sig')
                print('文件保存成功：', result_file)
                if not json_data['data']['has_more']:
                    break
                # 判断终止条件
                next_cursor = json_data['data']['cursor']
                page += 1
        df = pd.read_csv('小红书评论_{}.csv'.format(now))
        df.to_excel('xhs_1.xlsx', index=False)
        os.remove('小红书评论_{}.csv'.format(now))
    #情感分析
    def process_xiaohongshu_data(self):
        # 检查并处理已存在的文件
        if os.path.exists("话题笔记.csv"):
            pd.read_csv("话题笔记.csv").to_excel("note_1.xlsx", index=False)
            os.remove("话题笔记.csv")
        else:
            pass
        # 读取Excel文件并进行处理
        df = pd.read_excel('note_1.xlsx')
        df.sort_values(by='笔记发布时间', ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df.to_excel('note_1.xlsx', index=False)
        df = pd.read_excel('note_1.xlsx')

        def read_words_from_file(file_path):
            with open(file_path, 'r', encoding="utf-8") as file:
                words = file.read().strip()
                if words.endswith(','):
                    words = words[:-1]
                return words

        negative_words = read_words_from_file("negative_words.txt")
        positive_words = read_words_from_file("positive_words.txt")
        pass_words = read_words_from_file("pass_words.txt")

        # 情感分析函数
        def get_sentiment_gpt(text):
            try:
                if not text.strip():  # 如果文本为空白
                    return "中性"  # 直接返回"中性"
                s = SnowNLP(text)
                sentiment = s.sentiments
                if sentiment > 0.9:
                    return "积极"
                elif sentiment < 0.8:
                    return "消极"
                else:
                    return "中性"
                if any(word in text for word in negative_words):
                    return '消极'
                if any(word in text for word in positive_words):
                    return '积极'
                if any(word in text for word in pass_words):
                    return '中性'
            except:
                return '消极'

        # 对数据进行情感分析和筛选

        df['笔记内容'] = df['笔记内容'].fillna('').astype(str)
        df = df.drop_duplicates(subset=['笔记内容'])
        df['情感倾向_gpt'] = df['笔记内容'].apply(get_sentiment_gpt)
        df['人工情感分析'] = ''  # 添加新的空列

        negative_df = df[df['情感倾向_gpt'] == '消极']
        # 在排序时直接使用replace函数，而不创建新列
        negative_df = negative_df.sort_values(by='情感倾向_gpt', key=lambda col: col.replace({'消极': 0, '中性': 1, '积极': 2}),
                                              ascending=True)
        negative_df['人工情感分析'] = ''  # 添加新的空列

        df.to_excel('note_3.xlsx', index=False)
        negative_df.to_excel('note_2.xlsx', index=False)

        # 读取原始数据
        df = pd.read_excel('note_3.xlsx')

        # 筛选出非消极的情感倾向
        df_positive_neutral = df[df['情感倾向_gpt'] != '消极']

        # 将筛选结果写入新的Excel文件
        df_positive_neutral.to_excel('note_4.xlsx', index=False)

        # 读取新的Excel文件
        df_new = pd.read_excel('note_4.xlsx')

        # 清除指定的列的内容
        df_new[['情感倾向_gpt', '人工情感分析']] = None

        # 将结果写回新的Excel文件
        df_new.to_excel('note_4.xlsx', index=False)

        def get_sentiment_gpt(text):
            url = "http://159.138.108.97/newBear/gptAi.php?token=aSyBWgXYokswSNWzdYmq"
            headers = {"Content-Type": "application/json"}
            emotion = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"请从情感分析的角度分析直接给出这段文本是0到1哪个数，越靠近0越消极，越靠近1越积极: {text}"}
            ]
            response = requests.post(url, headers=headers, data=json.dumps(emotion))
            answer = response.json()
            sentiment_text = answer['answer']

            # 提取文本中的数字
            numbers = re.findall(r'\d+(?:\.\d+)?', sentiment_text)
            if not numbers:
                sentiment = 0.65
            elif len(numbers) == 1:
                sentiment = float(numbers[0])
            else:
                sentiment = round(np.mean([float(num) for num in numbers]), 2)

            # 根据得分判断情感倾向
            if sentiment > 0.7:
                sentiment = "积极"
            elif sentiment < 0.6:
                sentiment = "消极"
            else:
                sentiment = "中性"

            return sentiment

        # 读取新的Excel文件
        df_new = pd.read_excel('note_4.xlsx')

        # 对 '笔记内容' 列应用 get_sentiment_gpt 函数
        df_new['情感倾向_gpt'] = df_new['笔记内容'].apply(get_sentiment_gpt)

        df_new.to_excel('note_4.xlsx', index=False)

        df = pd.read_excel('note_4.xlsx', engine='openpyxl')

        # 找出情感倾向_gpt列中带有“消极”字样的行
        negative_rows = df[df['情感倾向_gpt'] == '消极']

        # 读取已存在的Excel文件
        df_existing = pd.read_excel('note_2.xlsx', engine='openpyxl')

        # 将negative_rows附加到已存在的数据帧末尾
        df_combined = pd.concat([df_existing, negative_rows], ignore_index=True)

        # 将结果写回Excel文件
        df_combined.to_excel('note_2.xlsx', index=False)
        df = pd.read_excel('note_2.xlsx')
        df.sort_values(by='笔记发布时间', ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df.to_excel('note_2.xlsx', index=False)


        # 使用Selenium模拟浏览器操作，并生成文本和截图
        dd = pd.read_excel('note_2.xlsx')

        # 遍历每一行
        for index, row in dd.iterrows():
            keyword_1 = row['笔记标题']
            if pd.isnull(keyword_1):
                keyword_1 = "标题太长"

            key = row["笔记关键词"]
            xhs_content = row['笔记内容']
            link = row['笔记链接']

            # 设置Chrome浏览器的选项
            options = Options()
            options.add_argument("--disable-blink-features=AutomationControlled")

            # 初始化webdriver
            driver = webdriver.Chrome(options=options)

            try:
                # 访问链接
                driver.get(url=link)

                # 设置显式等待，等待页面某个元素完全加载，这里以页面的body元素为例
                WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                time.sleep(3)
                # 保存截图
                driver.save_screenshot(f'screenshot_{index + 1}.png')
            except Exception as e:
                print(f"处理链接较慢，请稍等: {link}")
                print(str(e))
            finally:
                # 关闭浏览器
                driver.quit()

            # 随机等待3到5秒
            time.sleep(random.uniform(3, 5))

            url = "http://159.138.108.97/newBear/gptAi.php?token=aSyBWgXYokswSNWzdYmq"
            headers = {"Content-Type": "application/json"}
            messages = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"请为以下内容生成一个总结：\n{xhs_content}"}
            ]
            response = requests.post(url, headers=headers, data=json.dumps(messages))
            output_dict = response.json()
            summary_1 = output_dict['answer']


            with open(f'fu_{index}.txt', 'w', encoding='utf-8') as f:
                print(f'当前笔记的关键词为：{key}\n', file=f)
                print(f'标题：{", ".join(keyword_1)}\n', file=f)
                print(f'{summary_1}\n', file=f)


        xhs_1 = pd.read_excel('xhs_1.xlsx')
        note_2 = pd.read_excel('note_2.xlsx')
        xhs_1['ID'] = xhs_1['笔记链接'].apply(lambda x: x.split('/')[-1])
        note_2['ID'] = note_2['笔记链接'].apply(lambda x: x.split('/')[-1])

        common_ids = set(xhs_1['ID']).intersection(set(note_2['ID']))
        xhs_2 = pd.DataFrame()
        for id in note_2['ID']:
            if id in common_ids:
                xhs_2 = pd.concat([xhs_2, xhs_1[xhs_1['ID'] == id]], ignore_index=True)
        xhs_2 = xhs_2.drop(columns=['ID'])
        xhs_2.to_excel('xhs_2.xlsx', index=False)
        # 处理小红书的评论数据，并进行情感分析和文本摘要
        df = pd.read_excel('xhs_2.xlsx')
        df['评论内容'] = df['评论内容'].astype(str)
        df = df.groupby('笔记链接')['评论内容'].apply(' '.join).reset_index()
        for index, row in df.iterrows():
            xhs_content = row['评论内容']
            url = "http://159.138.108.97/newBear/gptAi.php?token=aSyBWgXYokswSNWzdYmq"
            headers = {"Content-Type": "application/json"}
            messages_1 = [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": f"请为以下内容生成一个总结：\n{xhs_content}"}
            ]
            output_dict = requests.post(url, headers=headers, data=json.dumps(messages_1))
            summary_2 = output_dict.json()
            if summary_2 == None:
                summary = "该话题下评论内容较多导致无法正常总结内容，可通过笔记链接自主访问该话题下评论内容"
            else:
                summary = summary_2['answer']



            s = SnowNLP(xhs_content)
            sentiment_1 = s.sentiments
            if sentiment_1 > 0.6:
                sentiment = '积极'
            elif sentiment_1 < 0.4:
                sentiment = '消极'
            else:
                sentiment = '中性'

            with open(f'fu_{index}.txt', 'a', encoding='utf-8') as f:
                print(f'评论区：{summary}\n', file=f)
                print(f'评论区主要情感：{sentiment}\n', file=f)

        # 处理笔记数据，并生成文本文件
        dg = pd.read_excel('note_2.xlsx')
        for index, row in dg.iterrows():
            keyword_1 = row['笔记标题']
            if pd.isnull(keyword_1):
                keyword_1 = "标题太长"
            post_time = row['笔记发布时间']
            link = row['笔记链接']

            with open(f'fu_{index}.txt', 'a', encoding='utf-8') as f:
                print(f'发帖时间：{post_time}\n', file=f)
                print(f'链接：{link}\n', file=f)

        # 将生成的文本和截图整合到Word文档中
        all_files = os.listdir()
        fu_files = sorted(
            [f for f in all_files if re.match(r'fu_\d+\.txt', f)],
            key=lambda x: int(re.search(r'(\d+)', x).group())
        )
        screenshot_files = sorted(
            [f for f in all_files if re.match(r'screenshot_\d+\.png', f)],
            key=lambda x: int(re.search(r'(\d+)', x).group())
        )
        doc = Document()
        max_files = max(len(fu_files), len(screenshot_files))
        for i in range(max_files):
            if i < len(screenshot_files):
                doc.add_picture(screenshot_files[i], width=Inches(7))
            if i < len(fu_files):
                with open(fu_files[i], 'r', encoding='utf-8') as file:
                    text = file.read()
                doc.add_paragraph(text)

        doc.save('小红书负面情绪测评.docx')
    #发送邮件
    def send_email_with_attachments(self):
        def attach_files(msg, file_type):
            files = glob.glob(f"*.{file_type}")
            for file in files:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(open(file, 'rb').read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', "attachment", filename=os.path.basename(file))
                msg.attach(part)

        from_addr = "2447956781@qq.com"
        password = "wmdcrvzkivilebca"
        emails = self.controller.pages[PageTwo].get_emails()

        # 创建邮件对象
        msg = MIMEMultipart()
        msg['From'] = from_addr
        msg['Subject'] = "小红书负面情绪测评"

        # 添加.docx文件作为附件
        attach_files(msg, "docx")

        # 添加.xlsx文件作为附件
        attach_files(msg, "xlsx")

        server = smtplib.SMTP('smtp.qq.com', 587)
        server.starttls()
        server.login(from_addr, password)

        for to_addr in emails:
            msg['To'] = to_addr
            text = msg.as_string()
            server.sendmail(from_addr, to_addr, text)

        server.quit()

    def stop_countdown(self):
        self.is_running = True

    def clear_times(self):
        for time_entry in self.times:
            time_entry.delete(0, tk.END)

class GUIApp(tk.Tk):
    def __init__(self, *args, **kwargs):
        tk.Tk.__init__(self, *args, **kwargs)

        container = tk.Frame(self)
        container.pack(side="top", fill="both", expand=True)

        self.pages = {}

        for Page in (PageOne, PageTwo, PageThree):
            page = Page(container, self)
            self.pages[Page] = page
            page.grid(row=0, column=0, sticky="nsew")

        self.show_page(PageOne)

    def show_page(self, Page):
        page = self.pages[Page]
        page.tkraise()

if __name__ == "__main__":
    app = GUIApp()
    app.mainloop()
